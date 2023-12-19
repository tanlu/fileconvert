# txt 和 docx  doc 之间的转换
from docx import Document
from flask import jsonify

from fileutils import get_file_name_no_extension

TYPE_DOC = '.doc'
TYPE_DOCX = '.docx'
TYPE_TXT = '.txt'


# txt - docx
def txt_to_docx(file_path, file_name):
    txt_full_path = file_path + file_name
    # 打开txt文件进行读取
    with open(txt_full_path, 'r', encoding='utf-8') as txt:
        content = txt.read()

    # 创建一个新的docx文档
    doc = Document()

    # 将txt内容添加到docx文档中
    doc.add_paragraph(content)
    # 构建docx 全路径
    fname = get_file_name_no_extension(file_name)
    docx_full_path = file_path + fname + TYPE_DOCX
    # 保存docx文档
    doc.save(docx_full_path)
    re = {"file": docx_full_path, "success": True}
    return jsonify(re)


# txt - docx
def txt_to_doc(file_path, file_name):
    txt_full_path = file_path + file_name
    # 打开txt文件进行读取
    with open(txt_full_path, 'r', encoding='utf-8') as txt:
        content = txt.read()

    # 创建一个新的docx文档
    doc = Document()
    # 将txt内容添加到docx文档中
    doc.add_paragraph(content)
    # 构建docx 全路径
    fname = get_file_name_no_extension(file_name)
    docx_full_path = file_path + fname + TYPE_DOC
    # 保存docx文档
    doc.save(docx_full_path)
    re = {"file": docx_full_path, "success": True}
    return jsonify(re)


#  docx/doc -> txt
def docx_to_txt(file_path, file_name):
    # 打开docx文件进行读取
    docx_full_path = file_path + file_name
    doc = Document(docx_full_path)

    # 提取文本内容
    content = []
    for paragraph in doc.paragraphs:
        content.append(paragraph.text)
    # 构建文本全路径
    fname = get_file_name_no_extension(file_name)
    txt_full_path = file_path + fname + TYPE_TXT
    # 将文本内容写入txt文件
    with open(txt_full_path, 'w', encoding='utf-8') as txt:
        txt.write('\n'.join(content))
    re = {"file": docx_full_path, "success": True}
    return jsonify(re)
